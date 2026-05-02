# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import numpy as np

# 尝试导入PDF转换库
try:
    from docx2pdf import convert as _docx2pdf_convert
except:
    _docx2pdf_convert = None

try:
    import win32com.client as _win32com
except:
    _win32com = None

import subprocess

# Set Chinese font
plt.rcParams['font.sans-serif'] = ['STSong']
plt.rcParams['axes.unicode_minus'] = False

# 使用相对路径配置
script_dir = os.path.dirname(os.path.abspath(__file__))
base_dir = os.path.dirname(script_dir)
xlsx_path_shared = os.path.join(base_dir, '..', 'data', '寄递时限对标分析指标大全.xlsx')
xlsx_path_local = os.path.join(base_dir, '输入数据', '接口文档-0401(1).xlsx')

if os.path.exists(xlsx_path_shared):
    xlsx_path = xlsx_path_shared
    print('使用共享数据文件：' + xlsx_path)
elif os.path.exists(xlsx_path_local):
    xlsx_path = xlsx_path_local
    print('使用本地数据文件：' + xlsx_path)
else:
    xlsx_path = xlsx_path_shared
    print('使用默认数据文件：' + xlsx_path)
output_dir = script_dir

# 读取各sheet数据
def read_sheets():
    # 线路表
    df_raw = pd.read_excel(xlsx_path, sheet_name='线路表', header=None)
    cols_route = df_raw.iloc[1].tolist()
    df_route = df_raw.iloc[2:].copy()
    df_route.columns = cols_route
    df_route.columns = [str(c).strip() for c in df_route.columns]
    
    # 进口机构表
    df_inst = pd.read_excel(xlsx_path, sheet_name='进口机构表', header=0)
    df_inst.columns = [str(c).strip() for c in df_inst.columns]
    
    # 48小时表
    df_48h = pd.read_excel(xlsx_path, sheet_name='进口、出口处理及投递48小时', header=0)
    df_48h['机构名称'] = df_48h['机构名称'].astype(str)
    df_48h['环节标识'] = df_48h['环节标识'].astype(str)
    df_48h['到达分拣离开标识'] = df_48h['到达分拣离开标识'].astype(str)
    df_48h['线路名'] = df_48h['线路名'].astype(str)
    
    return df_route, df_inst, df_48h

def convert_to_pdf(docx_path):
    try:
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'

        if _docx2pdf_convert is not None:
            try:
                _docx2pdf_convert(docx_path, pdf_path)
                print('PDF文档已保存: ' + pdf_path)
                return pdf_path
            except Exception as e1:
                print('docx2pdf转换失败: ' + str(e1))

        if _win32com is not None:
            try:
                word = _win32com.client.Dispatch('Word.Application')
                word.Visible = False
                word_doc = word.Documents.Open(docx_path)
                word_doc.SaveAs(pdf_path, FileFormat=17)
                word_doc.Close()
                word.Quit()
                print('PDF文档已保存: ' + pdf_path)
                return pdf_path
            except Exception as e2:
                print('win32com转换失败: ' + str(e2))

        try:
            lo_paths = ['libreoffice', 'soffice',
                        os.path.expanduser('~/.libreoffice/LibreOffice.app/Contents/MacOS/soffice'),
                        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                        os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(base_dir))), '.libreoffice', 'LibreOffice.app', 'Contents', 'MacOS', 'soffice')]
            lo_cmd = None
            for p in lo_paths:
                if os.path.isabs(p):
                    if os.path.exists(p):
                        lo_cmd = p
                        break
                else:
                    from shutil import which
                    if which(p):
                        lo_cmd = p
                        break
            if not lo_cmd:
                raise FileNotFoundError('LibreOffice未安装')
            result = subprocess.run(
                [lo_cmd, '--headless', '--convert-to', 'pdf', '--outdir',
                 os.path.dirname(pdf_path), docx_path],
                capture_output=True, text=True, timeout=60
            )
            if os.path.exists(pdf_path):
                print('PDF文档已保存（LibreOffice）: ' + pdf_path)
                return pdf_path
            else:
                print('LibreOffice转换失败: ' + result.stderr)
        except FileNotFoundError:
            print('LibreOffice未安装，跳过PDF转换')
        except Exception as e3:
            print('LibreOffice转换异常: ' + str(e3))

        print('PDF转换失败，仅生成Word文档')
        return None
    except Exception as e:
        print('PDF转换失败: ' + str(e))
        return None

def create_48h_chart(df_48h, inst_name, output_dir):
    """生成48小时到达/分拣/离开分布图"""
    inst_data = df_48h[(df_48h['机构名称'].str.contains(inst_name.replace(' ', ''), na=False)) & (df_48h.get('环节标识', '') == '进口')]
    if inst_data.empty:
        return None
    
    # 提取0-23小时的列
    hour_cols = []
    for c in df_48h.columns:
        try:
            if isinstance(c, (int, float)) and 0 <= float(c) <= 23:
                hour_cols.append(c)
        except:
            pass
    hour_cols = sorted(hour_cols, key=lambda x: int(float(x)))
    
    if not hour_cols:
        return None
    
    # 提取各标识的数据
    charts = {}
    for label in ['到达', '分拣', '离开']:
        label_data = inst_data[inst_data['到达分拣离开标识'].str.contains(label, na=False)]
        if not label_data.empty:
            row = label_data.iloc[0]
            values = []
            for c in hour_cols:
                try:
                    val = float(row[c]) if pd.notna(row[c]) else 0
                    values.append(val)
                except:
                    values.append(0)
            charts[label] = values
    
    if not charts:
        return None
    
    # 绘制组合图
    fig, ax = plt.subplots(figsize=(12, 6))
    x = range(len(hour_cols))
    colors = {'到达': '#1f77b4', '分拣': '#ff7f0e', '离开': '#2ca02c'}
    markers_map = {'到达': 'o', '分拣': 's', '离开': '^'}
    
    for label, values in charts.items():
        color = colors.get(label, '#999999')
        marker = markers_map.get(label, 'o')
        ax.plot(x, values, marker=marker, label=label, color=color, linewidth=2, markersize=6)
    
    ax.set_xlabel('Hour', fontsize=12)
    ax.set_ylabel('Ratio', fontsize=12)
    ax.set_title(inst_name + ' - 48h Distribution', fontsize=14, fontweight='bold')
    ax.legend(prop={'size': 10})
    ax.set_xticks(x)
    ax.set_xticklabels([str(int(c)) for c in hour_cols], rotation=45)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, '48h_{}_综合分布图.png'.format(inst_name))
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    # 波峰分析
    peak_info = {}
    for label, values in charts.items():
        if values and max(values) > 0:
            max_idx = np.argmax(values)
            peak_info[label] = {'hour': str(int(hour_cols[max_idx])), 'value': values[max_idx] * 100}
    
    return chart_path, peak_info

def create_48h_chart_from_data(inst_data_48h, inst_name, output_dir):
    """从已有的48小时数据生成图表"""
    # 提取0-23小时的列
    hour_cols = []
    for c in inst_data_48h.columns:
        try:
            if isinstance(c, (int, float)) and 0 <= float(c) <= 23:
                hour_cols.append(c)
        except:
            pass
    hour_cols = sorted(hour_cols, key=lambda x: int(float(x)))
    
    if not hour_cols:
        return None
    
    # 提取各标识的数据
    charts = {}
    for label in ['到达', '分拣', '离开']:
        label_data = inst_data_48h[inst_data_48h['到达分拣离开标识'].str.contains(label, na=False)]
        if not label_data.empty:
            row = label_data.iloc[0]
            values = []
            for c in hour_cols:
                try:
                    val = float(row[c]) if pd.notna(row[c]) else 0
                    values.append(val)
                except:
                    values.append(0)
            charts[label] = values
    
    if not charts:
        return None
    
    # 绘制组合图
    fig, ax = plt.subplots(figsize=(12, 6))
    x = range(len(hour_cols))
    colors = {'到达': '#1f77b4', '分拣': '#ff7f0e', '离开': '#2ca02c'}
    markers_map = {'到达': 'o', '分拣': 's', '离开': '^'}
    
    for label, values in charts.items():
        color = colors.get(label, '#999999')
        marker = markers_map.get(label, 'o')
        ax.plot(x, values, marker=marker, label=label, color=color, linewidth=2, markersize=6)
    
    ax.set_xlabel('Hour', fontsize=12)
    ax.set_ylabel('Ratio', fontsize=12)
    ax.set_title(inst_name + ' - 48h Import Distribution', fontsize=14, fontweight='bold')
    ax.legend(prop={'size': 10})
    ax.set_xticks(x)
    ax.set_xticklabels([str(int(c)) for c in hour_cols], rotation=45)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, '48h_{}_进口综合分布图.png'.format(inst_name))
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    # 波峰分析
    peak_info = {}
    for label, values in charts.items():
        if values and max(values) > 0:
            max_idx = np.argmax(values)
            peak_info[label] = {'hour': str(int(hour_cols[max_idx])), 'value': values[max_idx] * 100}
    
    return chart_path, peak_info

def generate_import_48h_analysis(charts_data, peak_info, inst_name, process_time, sort_time, wait_time, sample):
    """生成完整的48小时分布分析文本（大模型风格）"""
    analysis_parts = []
    
    # 总体概述
    analysis_parts.append('针对{}进口环节的48小时到达/分拣/离开分布分析显示'.format(inst_name))
    
    # 到达环节分析
    if '到达' in charts_data:
        arrival_values = charts_data['到达']
        arrival_hours = [(h, v) for h, v in enumerate(arrival_values) if v > 0]
        if arrival_hours:
            peak_h = max(arrival_hours, key=lambda x: x[1])[0]
            peak_v = max(arrival_hours, key=lambda x: x[1])[1]
            main_hours = sorted([h for h, v in arrival_hours if v >= peak_v * 0.3])
            
            if len(main_hours) >= 2:
                analysis_parts.append('在到达环节，邮件主要集中在{}点至{}点之间到达，其中{}点为关键波峰时段，到达占比达到{:.1f}%'.format(
                    min(main_hours), max(main_hours), peak_h, peak_v * 100))
            else:
                analysis_parts.append('在到达环节，邮件高度集中在{}点到达，占比达到{:.1f}%，呈现明显的时段集中特征'.format(
                    peak_h, peak_v * 100))
    
    # 分拣环节分析
    if '分拣' in charts_data:
        sort_values = charts_data['分拣']
        sort_hours = [(h, v) for h, v in enumerate(sort_values) if v > 0]
        if sort_hours:
            peak_h = max(sort_hours, key=lambda x: x[1])[0]
            peak_v = max(sort_hours, key=lambda x: x[1])[1]
            analysis_parts.append('分拣环节的关键波峰出现在{}点，分拣占比为{:.1f}%，表明邮件在到达后需要{:.2f}小时左右完成分拣处理'.format(
                peak_h, peak_v * 100, sort_time if pd.notna(sort_time) else 0))
    
    # 离开环节分析
    if '离开' in charts_data:
        leave_values = charts_data['离开']
        leave_hours = [(h, v) for h, v in enumerate(leave_values) if v > 0]
        if leave_hours:
            peak_h = max(leave_hours, key=lambda x: x[1])[0]
            peak_v = max(leave_hours, key=lambda x: x[1])[1]
            analysis_parts.append('离开环节的关键波峰出现在{}点，离开占比为{:.1f}%，表明邮件分拣完成后需要等待发运，待发运时限为{:.2f}小时'.format(
                peak_h, peak_v * 100, wait_time if pd.notna(wait_time) else 0))
    
    # 综合问题诊断
    analysis_parts.append('综合分析发现，{}进口处理时限长达{:.2f}小时，主要问题在于分拣环节耗时{:.2f}小时，占整体处理时限的{:.1f}%，是进口处理的主要瓶颈；待发运环节耗时{:.2f}小时，占整体处理时限的{:.1f}%，说明发运班次安排不够合理，存在积压现象'.format(
        inst_name,
        process_time if pd.notna(process_time) else 0,
        sort_time if pd.notna(sort_time) else 0,
        (sort_time / process_time * 100) if pd.notna(sort_time) and pd.notna(process_time) and process_time > 0 else 0,
        wait_time if pd.notna(wait_time) else 0,
        (wait_time / process_time * 100) if pd.notna(wait_time) and pd.notna(process_time) and process_time > 0 else 0))
    
    # 优化建议
    analysis_parts.append('针对上述问题，建议优化分拣流程以提高分拣效率，增加分拣设备投入实现自动化分拣，优化发运班次安排减少待发运积压时间，调整到达时段分布避免集中到达导致处理拥堵')
    
    return '。'.join(analysis_parts) + '。'

def main():
    print('读取数据...')
    df_route, df_inst, df_48h = read_sheets()
    
    # 创建Word文档
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('滨州市惠州市 线路进口环节时限对标下钻分析报告', 0)
    for run in title.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(24)
        run.font.bold = True
    
    date_str = datetime.now().strftime('%Y年%m月%d日')
    date_para = doc.add_paragraph('分析日期: ' + date_str)
    for run in date_para.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
    
    # 一、全程时限对标
    heading1 = doc.add_heading('一、全程时限对标', level=1)
    for run in heading1.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    # 查找滨州市惠州市线路数据
    route_name = '滨州市惠州市'
    route_data = df_route[df_route['线路名'] == route_name]
    if not route_data.empty:
        row = route_data.iloc[0]
        
        # 提取关键数据
        sample_post = pd.to_numeric(row.get('特快-样本量（件）', 0), errors='coerce')
        time_post = pd.to_numeric(row.get('特快-全程时限（小时）', 0), errors='coerce')
        time_comp = pd.to_numeric(row.get('竞品-全程时限（小时）', 0), errors='coerce')
        
        trajectory_post = str(row.get('主要轨迹-特快', ''))
        trajectory_comp = str(row.get('主要轨迹-竞品', ''))
        
        # 计算经转城市数量
        post_cities = len([c for c in trajectory_post.split('-') if c.strip()]) if trajectory_post else 0
        comp_cities = len([c for c in trajectory_comp.split('-') if c.strip()]) if trajectory_comp else 0
        
        # 进口环节时限
        import_post = pd.to_numeric(row.get('特快-进口时限（小时）', 0), errors='coerce')
        import_comp = pd.to_numeric(row.get('竞品-进口时限（小时）', 0), errors='coerce')
        import_diff = pd.to_numeric(row.get('差值-进口时限（小时）', 0), errors='coerce')
        
        # 进口处理细分数据
        import_process_post = pd.to_numeric(row.get('特快-进口处理时间（小时）', 0), errors='coerce')
        import_process_comp = pd.to_numeric(row.get('竞品-进口处理时间（小时）', 0), errors='coerce')
        import_process_diff = pd.to_numeric(row.get('差值-进口处理时间（小时）', 0), errors='coerce')
        
        import_transport_post = pd.to_numeric(row.get('特快-进口运输时间（小时）', 0), errors='coerce')
        import_transport_comp = pd.to_numeric(row.get('竞品-进口运输时间（小时）', 0), errors='coerce')
        import_transport_diff = pd.to_numeric(row.get('差值-进口运输时间（小时）', 0), errors='coerce')
        
        import_process_count_post = pd.to_numeric(row.get('特快-进口处理次数（次）', 0), errors='coerce')
        import_process_count_comp = pd.to_numeric(row.get('竞品-进口处理次数（次）', 0), errors='coerce')
        
        import_city_count_post = pd.to_numeric(row.get('特快-进口城市数（个）', 0), errors='coerce')
        import_city_count_comp = pd.to_numeric(row.get('竞品-进口城市数（个）', 0), errors='coerce')
        
        # 生成分析文本
        if pd.notna(time_post) and pd.notna(time_comp):
            if time_post > time_comp:
                diff_text = '特快比竞品慢{:.1f}小时'.format(time_post - time_comp)
            else:
                diff_text = '特快比竞品快{:.1f}小时'.format(time_comp - time_post)
            
            if post_cities == comp_cities:
                traj_text = '特快及竞品主要轨迹相同，均为{}'.format(trajectory_post)
            elif post_cities > comp_cities:
                traj_text = '特快主要轨迹为{}，比竞品多经转{}个城市'.format(trajectory_post, post_cities - comp_cities)
            else:
                traj_text = '竞品主要轨迹为{}，比特快多经转{}个城市'.format(trajectory_comp, comp_cities - post_cities)
            
            full_para = doc.add_paragraph(
                '以滨州市至惠州市线路为例，{}，{}。'.format(traj_text, diff_text)
            )
            for run in full_para.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(12)
            
            # 在终端打印
            print('\n' + '=' * 70)
            print('滨州市至惠州市 线路全程时限对标分析')
            print('=' * 70)
            print('样本量：{:.0f}件'.format(sample_post if pd.notna(sample_post) else 0))
            print('特快全程时限：{:.1f}小时'.format(time_post if pd.notna(time_post) else 0))
            print('竞品全程时限：{:.1f}小时'.format(time_comp if pd.notna(time_comp) else 0))
            print('全程时限差异：{}'.format(diff_text))
    else:
        print('未找到滨州市惠州市线路数据')
        para = doc.add_paragraph('未找到滨州市至惠州市线路的完整数据。')
        for run in para.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
        return
    
    # 二、进口环节对标
    heading2 = doc.add_heading('二、进口环节对标', level=1)
    for run in heading2.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    if pd.notna(import_post) and pd.notna(import_comp):
        if import_post > import_comp:
            import_analysis = '进口环节特快慢于竞品{:.1f}小时，是需要重点优化的环节。'.format(import_post - import_comp)
        else:
            import_analysis = '进口环节特快优于竞品{:.1f}小时，是邮政的竞争优势环节。'.format(import_comp - import_post)
        
        para2 = doc.add_paragraph(
            '进口环节方面，特快进口时限为{:.1f}小时，竞品为{:.1f}小时。{}'.format(
                import_post, import_comp, import_analysis)
        )
        for run in para2.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
        
        print('\n进口环节分析：{}'.format(import_analysis))
        
        # 进口处理细分分析
        print('\n' + '-' * 70)
        print('进口环节细分对标：')
        print('-' * 70)
        
        if pd.notna(import_process_post) and pd.notna(import_process_comp):
            print('进口处理时间：特快{:.1f}小时，竞品{:.1f}小时，差值{:.1f}小时'.format(
                import_process_post, import_process_comp, import_process_diff if pd.notna(import_process_diff) else 0))
        
        if pd.notna(import_transport_post) and pd.notna(import_transport_comp):
            print('进口运输时间：特快{:.1f}小时，竞品{:.1f}小时，差值{:.1f}小时'.format(
                import_transport_post, import_transport_comp, import_transport_diff if pd.notna(import_transport_diff) else 0))
        
        if pd.notna(import_process_count_post) and pd.notna(import_process_count_comp):
            print('进口处理次数：特快{:.0f}次，竞品{:.0f}次'.format(
                import_process_count_post, import_process_count_comp))
        
        if pd.notna(import_city_count_post) and pd.notna(import_city_count_comp):
            print('进口城市数：特快{:.0f}个，竞品{:.0f}个'.format(
                import_city_count_post, import_city_count_comp))
    
    # 三、进口机构下钻分析
    heading3 = doc.add_heading('三、进口机构下钻分析', level=1)
    for run in heading3.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    # 获取end5进口机构
    route_inst = df_inst[df_inst['线路名'] == route_name].copy()
    if not route_inst.empty:
        for col in ['进口样本量', '进口处理时限', '进口分拣时限', '进口带发运时限']:
            if col in route_inst.columns:
                route_inst[col] = pd.to_numeric(route_inst[col], errors='coerce')
        route_inst = route_inst.sort_values(by=['进口处理时限', '进口样本量'], ascending=[False, False])
        end5 = route_inst.head(5)
        
        # end5表格
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ['进口机构名称', '样本量', '处理时限（小时）', '分拣时限（小时）', '待发运时限（小时）']
        for j, h in enumerate(headers):
            hdr_cells[j].text = h
            for paragraph in hdr_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
                    run.font.bold = True
        
        for _, r in end5.iterrows():
            row_cells = table.add_row().cells
            for j, h in enumerate(headers):
                val = r.get(h, '')
                if pd.isna(val):
                    val = '-'
                elif isinstance(val, float):
                    val = '{:.2f}'.format(val)
                row_cells[j].text = str(val)
                for paragraph in row_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
        
        doc.add_paragraph('')
        
        # 在终端打印end5机构数据
        print('\n' + '-' * 70)
        print('进口处理时限end5问题机构：')
        print('-' * 70)
        print('{:<20} {:>8} {:>12} {:>12} {:>12}'.format(
            '机构名称', '样本量', '处理时限', '分拣时限', '待发运时限'))
        print('-' * 70)
        for _, r in end5.iterrows():
            name = str(r.get('进口机构名称', ''))[:18]
            sample = '{:.0f}'.format(r.get('进口样本量', 0)) if pd.notna(r.get('进口样本量', 0)) else '-'
            process_time = '{:.2f}'.format(r.get('进口处理时限', 0)) if pd.notna(r.get('进口处理时限', 0)) else '-'
            sort_time = '{:.2f}'.format(r.get('进口分拣时限', 0)) if pd.notna(r.get('进口分拣时限', 0)) else '-'
            wait_time = '{:.2f}'.format(r.get('进口带发运时限', 0)) if pd.notna(r.get('进口带发运时限', 0)) else '-'
            print('{:<20} {:>8} {:>12} {:>12} {:>12}'.format(
                name, sample, process_time, sort_time, wait_time))
        
        # 四、问题机构48小时分布分析
        heading4 = doc.add_heading('四、问题机构48小时分布分析', level=1)
        for run in heading4.runs:
            run.font.name = '微软雅黑'
            run.font.bold = True
        
        # 针对最严重问题机构进行分析
        worst_inst = end5.iloc[0]
        worst_inst_name = str(worst_inst.get('进口机构名称', ''))
        
        if worst_inst_name:
            inst_heading = doc.add_heading('{} 下钻分析'.format(worst_inst_name), level=2)
            for run in inst_heading.runs:
                run.font.name = '微软雅黑'
                run.font.bold = True
            
            # 数据描述
            sample = worst_inst.get('进口样本量', 0)
            process_time = worst_inst.get('进口处理时限', 0)
            sort_time = worst_inst.get('进口分拣时限', 0)
            wait_time = worst_inst.get('进口带发运时限', 0)
            
            detail_parts = []
            if pd.notna(sample):
                detail_parts.append('样本量{:.0f}件'.format(sample))
            if pd.notna(process_time):
                detail_parts.append('处理时限{:.2f}小时'.format(process_time))
            if pd.notna(sort_time):
                detail_parts.append('分拣时限{:.2f}小时'.format(sort_time))
            if pd.notna(wait_time):
                detail_parts.append('待发运时限{:.2f}小时'.format(wait_time))
            
            detail_para = doc.add_paragraph('该机构进口处理环节数据：' + '，'.join(detail_parts) + '。')
            for run in detail_para.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(12)
            
            # 生成48小时分布图 - 需要匹配正确的机构名称
            # 进口机构在48小时表中的名称可能与进口机构表不同，需要查找匹配
            possible_names = [worst_inst_name, worst_inst_name.replace('进口', ''), worst_inst_name.replace('中心', '')]
            inst_data_48h = None
            matched_name = None
            
            for name in possible_names:
                inst_data_48h = df_48h[(df_48h['机构名称'].str.contains(name.replace(' ', ''), na=False)) & 
                                       (df_48h.get('环节标识', '') == '进口') &
                                       (df_48h['线路名'] == route_name)]
                if not inst_data_48h.empty:
                    matched_name = name
                    break
            
            # 如果未找到特定机构数据，尝试查找该线路的所有进口48小时数据
            if inst_data_48h is None or inst_data_48h.empty:
                inst_data_48h = df_48h[(df_48h.get('环节标识', '') == '进口') &
                                       (df_48h['线路名'] == route_name)]
                if not inst_data_48h.empty:
                    matched_name = str(inst_data_48h.iloc[0]['机构名称'])
            
            chart_result = None
            if inst_data_48h is not None and not inst_data_48h.empty:
                chart_result = create_48h_chart_from_data(inst_data_48h, matched_name or worst_inst_name, output_dir)
            
            if chart_result:
                chart_path, peak_info = chart_result
                if os.path.exists(chart_path):
                    doc.add_paragraph('')
                    doc.add_paragraph('{} 48小时到达/分拣/离开分布图：'.format(worst_inst_name))
                    doc.add_picture(chart_path, width=Inches(6))
                    
                    # 在终端显示分布数据表格
                    print('\n' + '=' * 70)
                    print('{} 48小时分布数据'.format(matched_name or worst_inst_name))
                    print('=' * 70)
                    
                    hour_cols_chart = sorted([c for c in df_48h.columns if isinstance(c, (int, float)) and 0 <= float(c) <= 23], key=lambda x: int(float(x)))
                    
                    charts_data = {}
                    for label in ['到达', '分拣', '离开']:
                        label_data = inst_data_48h[inst_data_48h['到达分拣离开标识'].str.contains(label, na=False)]
                        if not label_data.empty:
                            row = label_data.iloc[0]
                            values = []
                            for c in hour_cols_chart:
                                try:
                                    val = float(row[c]) if pd.notna(row[c]) else 0
                                    values.append(val)
                                except:
                                    values.append(0)
                            charts_data[label] = values
                    
                    # 打印完整句子形式的分布数据
                    print('\n【各环节48小时分布详情】')
                    print('-' * 70)
                    
                    for label, values in charts_data.items():
                        print('\n{}环节分布情况：'.format(label))
                        
                        # 找出有数据的时间点
                        active_hours = [(int(h), values[int(h)]) for h in hour_cols_chart if int(h) < len(values) and values[int(h)] > 0]
                        
                        if not active_hours:
                            print('  该环节暂无数据。')
                            continue
                        
                        # 生成完整句子描述
                        if len(active_hours) == 1:
                            h, v = active_hours[0]
                            print('  所有邮件均集中在{}点{}，占比{:.1f}%，呈现高度集中的分布特征。'.format(
                                h, label, v * 100))
                        elif len(active_hours) == 2:
                            h1, v1 = active_hours[0]
                            h2, v2 = active_hours[1]
                            print('  邮件主要分布在{}点和{}点两个时间段，其中{}点占比{:.1f}%，{}点占比{:.1f}%。'.format(
                                h1, h2, h1, v1 * 100, h2, v2 * 100))
                            if v1 > v2:
                                print('  {}点为关键波峰时段，占比最高，达到{:.1f}%。'.format(h1, v1 * 100))
                            else:
                                print('  {}点为关键波峰时段，占比最高，达到{:.1f}%。'.format(h2, v2 * 100))
                        else:
                            # 找出波峰
                            peak_h, peak_v = max(active_hours, key=lambda x: x[1])
                            total_active = sum(v for _, v in active_hours)
                            main_hours = sorted([h for h, v in active_hours if v >= peak_v * 0.5])
                            
                            print('  该环节邮件分布在多个时间段，主要集中在{}点至{}点之间。'.format(
                                min(main_hours), max(main_hours)))
                            print('  关键波峰出现在{}点，占比{:.1f}%，是邮件到达/处理/离开的最集中时段。'.format(
                                peak_h, peak_v * 100))
                            
                            if peak_v > 0.5:
                                print('  该时段邮件量占比超过50%，说明存在明显的时段集中现象，可能导致处理能力不足。')
                    
                    print('\n' + '-' * 70)
                    
                    print('\n[图表已保存]: ' + chart_path)
                    
                    # 波峰分析
                    if peak_info:
                        peak_parts = []
                        for label, info in peak_info.items():
                            peak_parts.append('{}关键波峰在{}点，占比{:.1f}%'.format(label, info['hour'], info['value']))
                        
                        analysis_para = doc.add_paragraph('波峰分析：' + '；'.join(peak_parts) + '。')
                        for run in analysis_para.runs:
                            run.font.name = '微软雅黑'
                            run.font.size = Pt(12)
                    
                    # 问题诊断
                    diagnosis_heading = doc.add_heading('问题诊断', level=3)
                    for run in diagnosis_heading.runs:
                        run.font.name = '微软雅黑'
                        run.font.bold = True
                    
                    diag_para = doc.add_paragraph('{}进口处理时限长达{:.2f}小时，主要原因分析：'.format(worst_inst_name, process_time if pd.notna(process_time) else 0))
                    for run in diag_para.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(12)
                    
                    diagnosis_parts = [
                        '分拣时限为{:.2f}小时，占处理时限的{:.1f}%，说明分拣环节耗时较长。'.format(
                            sort_time if pd.notna(sort_time) else 0,
                            (sort_time / process_time * 100) if pd.notna(sort_time) and pd.notna(process_time) and process_time > 0 else 0),
                        '待发运时限为{:.2f}小时，占处理时限的{:.1f}%，说明待发运环节存在积压。'.format(
                            wait_time if pd.notna(wait_time) else 0,
                            (wait_time / process_time * 100) if pd.notna(wait_time) and pd.notna(process_time) and process_time > 0 else 0),
                    ]
                    for part in diagnosis_parts:
                        p = doc.add_paragraph(part)
                        for run in p.runs:
                            run.font.name = '微软雅黑'
                            run.font.size = Pt(12)
                    
                    # 48小时分布完整分析（大模型生成）
                    analysis_heading = doc.add_heading('48小时分布完整分析', level=3)
                    for run in analysis_heading.runs:
                        run.font.name = '微软雅黑'
                        run.font.bold = True
                    
                    # 生成完整的48小时分布分析文本
                    full_analysis = generate_import_48h_analysis(charts_data, peak_info, worst_inst_name, 
                                                                 process_time, sort_time, wait_time, sample)
                    
                    analysis_para = doc.add_paragraph(full_analysis)
                    for run in analysis_para.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(12)
                    
                    # 在终端也打印完整分析
                    print('\n【48小时分布完整分析】')
                    print('-' * 70)
                    print(full_analysis)
                    print('-' * 70)
                    
                    # 优化建议
                    suggestion_heading = doc.add_heading('优化建议', level=3)
                    for run in suggestion_heading.runs:
                        run.font.name = '微软雅黑'
                        run.font.bold = True
                    
                    suggestions = [
                        '优化分拣流程，提高分拣效率，缩短分拣时限',
                        '增加分拣设备投入，实现自动化分拣',
                        '优化发运班次安排，减少待发运积压时间',
                        '调整到达时段分布，避免集中到达导致处理拥堵'
                    ]
                    for i, sug in enumerate(suggestions, 1):
                        p = doc.add_paragraph('{}. {}'.format(i, sug))
                        for run in p.runs:
                            run.font.name = '微软雅黑'
                            run.font.size = Pt(12)
    else:
        print('\n未找到滨州市惠州市线路的进口机构数据')
    
    # 五、分析总结
    heading5 = doc.add_heading('五、分析总结', level=1)
    for run in heading5.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    if pd.notna(time_post) and pd.notna(time_comp):
        summary_parts = [
            '滨州市至惠州市线路特快全程时限为{:.1f}小时，竞品为{:.1f}小时'.format(time_post, time_comp),
            '特快比竞品{}{:.1f}小时'.format('慢' if time_post > time_comp else '快', abs(time_post - time_comp))
        ]
        
        if pd.notna(import_post) and pd.notna(import_comp):
            summary_parts.append('进口环节特快为{:.1f}小时，竞品为{:.1f}小时'.format(import_post, import_comp))
            if import_post > import_comp:
                summary_parts.append('进口环节慢于竞品{:.1f}小时'.format(import_post - import_comp))
            else:
                summary_parts.append('进口环节优于竞品{:.1f}小时'.format(import_comp - import_post))
        
        summary_parts.append('需重点关注进口环节优化，提升整体时效表现。')
        
        summary = doc.add_paragraph('。'.join(summary_parts) + '。')
        for run in summary.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
    
    # 保存文档
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    output_path = os.path.join(output_dir, '线路进口环节时限分析报告_滨州市惠州市_{}.docx'.format(timestamp))
    doc.save(output_path)
    print('\nWord文档已保存: ' + output_path)
    
    # 显示48小时分布图路径
    chart_files = [f for f in os.listdir(output_dir) if f.startswith('48h_') and f.endswith('.png')]
    if chart_files:
        chart_path = os.path.join(output_dir, chart_files[0])
        print('\n48小时分布图: ' + chart_path)
    
    # 转换为PDF
    convert_to_pdf(output_path)
    
    print('\n分析完成！')

if __name__ == '__main__':
    main()
