# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

# 尝试导入 PDF 转换库
try:
    from docx2pdf import convert
except:
    convert = None

try:
    import win32com.client
except:
    win32com = None

import subprocess

# 使用相对路径配置
script_dir = os.path.dirname(os.path.abspath(__file__))
base_dir = os.path.dirname(script_dir)
xlsx_path_shared = os.path.join(base_dir, '..', 'data', '寄递时限对标分析指标大全.xlsx')
xlsx_path_local = os.path.join(base_dir, '输入数据', '接口文档-0401(1).xlsx')

if os.path.exists(xlsx_path_shared):
    xlsx_path = xlsx_path_shared
    print('使用共享数据文件：' + xlsx_path)
elif os.path.exists(xlsx_path_local):
    xlsx_path = xlsx_path_local
    print('使用本地数据文件：' + xlsx_path)
else:
    xlsx_path = xlsx_path_shared
    print('使用默认数据文件：' + xlsx_path)

output_dir = os.path.join(base_dir, '输出结果')
if not os.path.exists(output_dir):
    os.makedirs(output_dir, exist_ok=True)

# 读取线路表数据
def read_route_data():
    df_raw = pd.read_excel(xlsx_path, sheet_name='线路表', header=None)
    cols_route = df_raw.iloc[1].tolist()
    df_route = df_raw.iloc[2:].copy()
    df_route.columns = cols_route
    df_route.columns = [str(c).strip() for c in df_route.columns]
    return df_route

def convert_to_pdf(docx_path):
    try:
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'

        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            chinese_font_paths = [
                '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
                'C:\\Windows\\Fonts\\simhei.ttf',
                '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf'
            ]
            for fp in chinese_font_paths:
                if os.path.exists(fp):
                    try:
                        pdf.add_font('chinese', '', fp)
                        pdf.set_font('chinese', '', 14)
                        break
                    except Exception:
                        continue
            else:
                pdf.set_font('Helvetica', '', 14)

            doc = Document(docx_path)
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    pdf.ln(5)
                    continue
                style_name = para.style.name if para.style else ''
                if 'Heading 1' in style_name:
                    pdf.set_font('chinese', '', 16)
                    pdf.ln(8)
                elif 'Heading 2' in style_name:
                    pdf.set_font('chinese', '', 14)
                    pdf.ln(5)
                elif 'Heading 3' in style_name:
                    pdf.set_font('chinese', '', 12)
                    pdf.ln(3)
                else:
                    pdf.set_font('chinese', '', 11)
                pdf.multi_cell(0, 8, text)
                pdf.ln(2)

            try:
                for table in doc.tables:
                    pdf.ln(3)
                    pdf.set_font('chinese', '', 8)
                    for row_idx, row in enumerate(table.rows):
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                pdf.multi_cell(0, 5, cell_text)
                        pdf.ln(2)
                    pdf.ln(2)
            except Exception:
                pass

            pdf.output(pdf_path)
            print('PDF文档已保存（fpdf2）: ' + pdf_path)
            return pdf_path
        except Exception as e:
            print('fpdf2转换失败: ' + str(e))

        if convert is not None:
            try:
                convert(docx_path, pdf_path)
                print('PDF 文档已保存：' + pdf_path)
                return pdf_path
            except Exception as e1:
                print('docx2pdf 转换失败：' + str(e1))

        if win32com is not None:
            try:
                word = win32com.client.Dispatch('Word.Application')
                word.Visible = False
                word_doc = word.Documents.Open(docx_path)
                word_doc.SaveAs(pdf_path, FileFormat=17)
                word_doc.Close()
                word.Quit()
                print('PDF 文档已保存：' + pdf_path)
                return pdf_path
            except Exception as e2:
                print('win32com 转换失败：' + str(e2))

        try:
            lo_paths = ['libreoffice', 'soffice',
                        os.path.expanduser('~/.libreoffice/LibreOffice.app/Contents/MacOS/soffice'),
                        '/Applications/LibreOffice.app/Contents/MacOS/soffice']
            if 'base_dir' in dir():
                lo_paths.append(os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(base_dir))), '.libreoffice', 'LibreOffice.app', 'Contents', 'MacOS', 'soffice'))
            lo_cmd = None
            for p in lo_paths:
                if os.path.isabs(p):
                    if os.path.exists(p):
                        lo_cmd = p
                        break
                else:
                    from shutil import which
                    if which(p):
                        lo_cmd = p
                        break
            if not lo_cmd:
                raise FileNotFoundError('LibreOffice未安装')
            result = subprocess.run(
                [lo_cmd, '--headless', '--convert-to', 'pdf', '--outdir',
                 os.path.dirname(pdf_path), docx_path],
                capture_output=True, text=True, timeout=60
            )
            if os.path.exists(pdf_path):
                print('PDF 文档已保存（LibreOffice）：' + pdf_path)
                return pdf_path
            else:
                print('LibreOffice 转换失败：' + result.stderr)
        except FileNotFoundError:
            print('LibreOffice 未安装，跳过 PDF 转换')
        except Exception as e3:
            print('LibreOffice 转换异常：' + str(e3))

        print('PDF 转换失败，仅生成 Word 文档')
        return None
    except Exception as e:
        print('PDF 转换失败：' + str(e))
        return None

def generate_transit_analysis(route_name, row):
    """生成中转环节完整分析文本（大模型风格）"""
    analysis_parts = []
    
    # 提取关键数据
    time_post = pd.to_numeric(row.get('特快-全程时限（小时）', 0), errors='coerce')
    time_comp = pd.to_numeric(row.get('竞品-全程时限（小时）', 0), errors='coerce')
    
    transit_post = pd.to_numeric(row.get('特快-中转时限（小时）', 0), errors='coerce')
    transit_comp = pd.to_numeric(row.get('竞品-中转时限（小时）', 0), errors='coerce')
    transit_diff = pd.to_numeric(row.get('差值-中转时限（小时）', 0), errors='coerce')
    
    transit_cities_post = pd.to_numeric(row.get('特快-中转城市数', 0), errors='coerce')
    transit_cities_comp = pd.to_numeric(row.get('竞品-中转城市数', 0), errors='coerce')
    
    transit_orgs_post = pd.to_numeric(row.get('特快-中转机构数', 0), errors='coerce')
    transit_orgs_comp = pd.to_numeric(row.get('竞品-中转机构数', 0), errors='coerce')
    
    transit_distance_post = pd.to_numeric(row.get('特快-中转段里程（千米）', 0), errors='coerce')
    transit_distance_comp = pd.to_numeric(row.get('竞品-中转段里程（千米）', 0), errors='coerce')
    
    process_time_post = pd.to_numeric(row.get('特快-单次处理时长', 0), errors='coerce')
    process_time_comp = pd.to_numeric(row.get('竞品-单次处理时长', 0), errors='coerce')
    
    transport_time_post = pd.to_numeric(row.get('特快-运输时长', 0), errors='coerce')
    transport_time_comp = pd.to_numeric(row.get('竞品-运输时长', 0), errors='coerce')
    
    air_land_post = str(row.get('特快-航空陆运标识', ''))
    air_land_comp = str(row.get('竞品-航空陆运标识', ''))
    
    # 总体概述
    if pd.notna(time_post) and pd.notna(time_comp):
        if time_post > time_comp:
            diff_text = '特快比竞品慢{:.1f}小时'.format(time_post - time_comp)
        else:
            diff_text = '特快比竞品快{:.1f}小时'.format(time_comp - time_post)
        
        analysis_parts.append('{}，特快全程时限为{:.1f}小时，竞品为{:.1f}小时，{}'.format(
            route_name, time_post, time_comp, diff_text))
    
    # 中转环节分析
    if pd.notna(transit_post) and pd.notna(transit_comp):
        process_diff = None
        transport_diff = None
        cities_diff = None
        orgs_diff = None
        
        if pd.notna(transit_diff):
            if transit_diff > 0:
                transit_analysis = '其中中转环节特快慢于竞品{:.1f}小时'.format(transit_diff)
            else:
                transit_analysis = '其中中转环节特快优于竞品{:.1f}小时'.format(abs(transit_diff))
            
            analysis_parts.append(transit_analysis)
            
            # 细分环节分析
            if pd.notna(process_time_post) and pd.notna(process_time_comp):
                process_diff = process_time_post - process_time_comp
                if process_diff > 0:
                    analysis_parts.append('单次处理时长特快为{:.1f}小时，竞品为{:.1f}小时，特快慢于竞品{:.1f}小时'.format(
                        process_time_post, process_time_comp, process_diff))
                else:
                    analysis_parts.append('单次处理时长特快为{:.1f}小时，竞品为{:.1f}小时，特快优于竞品{:.1f}小时'.format(
                        process_time_post, process_time_comp, abs(process_diff)))
            
            if pd.notna(transport_time_post) and pd.notna(transport_time_comp):
                transport_diff = transport_time_post - transport_time_comp
                if transport_diff > 0:
                    analysis_parts.append('运输时长特快为{:.1f}小时，竞品为{:.1f}小时，特快慢于竞品{:.1f}小时'.format(
                        transport_time_post, transport_time_comp, transport_diff))
                else:
                    analysis_parts.append('运输时长特快为{:.1f}小时，竞品为{:.1f}小时，特快优于竞品{:.1f}小时'.format(
                        transport_time_post, transport_time_comp, abs(transport_diff)))
            
            # 中转城市数对比
            if pd.notna(transit_cities_post) and pd.notna(transit_cities_comp):
                cities_diff = transit_cities_post - transit_cities_comp
                if cities_diff > 0:
                    analysis_parts.append('中转城市数特快为{:.0f}个，竞品为{:.0f}个，特快比竞品多{:.0f}个城市'.format(
                        transit_cities_post, transit_cities_comp, cities_diff))
                elif cities_diff < 0:
                    analysis_parts.append('中转城市数特快为{:.0f}个，竞品为{:.0f}个，特快比竞品少{:.0f}个城市'.format(
                        transit_cities_post, transit_cities_comp, abs(cities_diff)))
                else:
                    analysis_parts.append('中转城市数特快与竞品相同，均为{:.0f}个城市'.format(transit_cities_post))
            
            # 中转机构数对比
            if pd.notna(transit_orgs_post) and pd.notna(transit_orgs_comp):
                orgs_diff = transit_orgs_post - transit_orgs_comp
                if orgs_diff > 0:
                    analysis_parts.append('中转机构数特快为{:.0f}个，竞品为{:.0f}个，特快比竞品多{:.0f}个机构'.format(
                        transit_orgs_post, transit_orgs_comp, orgs_diff))
                elif orgs_diff < 0:
                    analysis_parts.append('中转机构数特快为{:.0f}个，竞品为{:.0f}个，特快比竞品少{:.0f}个机构'.format(
                        transit_orgs_post, transit_orgs_comp, abs(orgs_diff)))
            
            # 中转里程对比
            if pd.notna(transit_distance_post) and pd.notna(transit_distance_comp):
                distance_diff = transit_distance_post - transit_distance_comp
                if distance_diff > 0:
                    analysis_parts.append('中转段里程特快为{:.0f}千米，竞品为{:.0f}千米，特快比竞品长{:.0f}千米'.format(
                        transit_distance_post, transit_distance_comp, distance_diff))
                elif distance_diff < 0:
                    analysis_parts.append('中转段里程特快为{:.0f}千米，竞品为{:.0f}千米，特快比竞品短{:.0f}千米'.format(
                        transit_distance_post, transit_distance_comp, abs(distance_diff)))
            
            # 航空陆运标识
            if air_land_post or air_land_comp:
                analysis_parts.append('运输方式：特快为{}，竞品为{}'.format(
                    air_land_post if air_land_post else '未知', 
                    air_land_comp if air_land_comp else '未知'))
            
            # 问题诊断
            if transit_diff > 0:
                analysis_parts.append('综合分析发现，中转环节时限差异的主要原因为：')
                
                if process_diff is not None and process_diff > 0:
                    analysis_parts.append('单次处理时长差异{:.1f}小时，说明特快在中转处理效率上存在不足'.format(process_diff))
                
                if transport_diff is not None and transport_diff > 0:
                    analysis_parts.append('运输时长差异{:.1f}小时，说明特快在运输组织上存在优化空间'.format(transport_diff))
                
                if cities_diff is not None and cities_diff > 0:
                    analysis_parts.append('中转城市数多{:.0f}个，导致邮件经转次数增加，影响整体时限'.format(cities_diff))
                
                if orgs_diff is not None and orgs_diff > 0:
                    analysis_parts.append('中转机构数多{:.0f}个，增加了处理环节和时间成本'.format(orgs_diff))
                
                # 优化建议
                analysis_parts.append('针对上述问题，建议：优化中转网络布局，减少不必要的中转城市；提升中转机构处理效率，缩短单次处理时长；优化运输组织，提高运输时效；根据邮件量合理配置航空陆运资源')
    
    return '。'.join(analysis_parts) + '。'

def main():
    print('读取数据...')
    df_route = read_route_data()
    
    # 创建 Word 文档
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('线路中转环节时限对标下钻分析报告', 0)
    for run in title.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(24)
        run.font.bold = True
    
    date_str = datetime.now().strftime('%Y年%m月%d日')
    date_para = doc.add_paragraph('分析日期：' + date_str)
    for run in date_para.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
    
    # 遍历所有线路
    unique_routes = df_route['线路名'].unique()
    
    for route_name in unique_routes:
        route_data = df_route[df_route['线路名'] == route_name]
        if not route_data.empty:
            row = route_data.iloc[0]
            
            # 提取关键数据
            sample_post = pd.to_numeric(row.get('特快-样本量（件）', 0), errors='coerce')
            time_post = pd.to_numeric(row.get('特快-全程时限（小时）', 0), errors='coerce')
            time_comp = pd.to_numeric(row.get('竞品-全程时限（小时）', 0), errors='coerce')
            
            # 中转环节数据
            transit_post = pd.to_numeric(row.get('特快-中转时限（小时）', 0), errors='coerce')
            transit_comp = pd.to_numeric(row.get('竞品-中转时限（小时）', 0), errors='coerce')
            transit_diff = pd.to_numeric(row.get('差值-中转时限（小时）', 0), errors='coerce')
            
            # 在终端打印
            print('\n' + '=' * 70)
            print('{} 线路中转环节时限对标分析'.format(route_name))
            print('=' * 70)
            print('样本量：{:.0f}件'.format(sample_post if pd.notna(sample_post) else 0))
            print('特快全程时限：{:.1f}小时'.format(time_post if pd.notna(time_post) else 0))
            print('竞品全程时限：{:.1f}小时'.format(time_comp if pd.notna(time_comp) else 0))
            
            if pd.notna(transit_post) and pd.notna(transit_comp):
                print('\n中转环节时限：')
                print('  特快中转时限：{:.1f}小时'.format(transit_post))
                print('  竞品中转时限：{:.1f}小时'.format(transit_comp))
                if pd.notna(transit_diff):
                    if transit_diff > 0:
                        print('  差异：特快慢于竞品{:.1f}小时'.format(transit_diff))
                    else:
                        print('  差异：特快优于竞品{:.1f}小时'.format(abs(transit_diff)))
            
            # 生成完整分析文本
            full_analysis = generate_transit_analysis(route_name, row)
            
            # 添加到文档
            analysis_heading = doc.add_heading('{} 中转环节时限对标分析'.format(route_name), level=1)
            for run in analysis_heading.runs:
                run.font.name = '微软雅黑'
                run.font.bold = True
            
            analysis_para = doc.add_paragraph(full_analysis)
            for run in analysis_para.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(12)
            
            # 在终端打印完整分析
            print('\n' + '-' * 70)
            print('【中转环节完整分析】')
            print('-' * 70)
            print(full_analysis)
            print('-' * 70)
    
    # 二、优化建议
    heading2 = doc.add_heading('优化建议', level=1)
    for run in heading2.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    suggestions = [
        '优化中转网络布局，减少不必要的中转城市，缩短中转路径',
        '提升中转机构处理效率，通过自动化设备缩短单次处理时长',
        '优化运输组织，提高运输时效，减少在途时间',
        '根据邮件量和时限要求合理配置航空陆运资源',
        '加强与竞品对标学习，借鉴先进管理经验'
    ]
    for i, sug in enumerate(suggestions, 1):
        p = doc.add_paragraph('{}. {}'.format(i, sug))
        for run in p.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
    
    # 保存文档
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    output_path = os.path.join(output_dir, '线路中转环节时限对标分析报告_{}.docx'.format(timestamp))
    doc.save(output_path)
    print('\nWord 文档已保存：' + output_path)
    
    # 转换为 PDF
    convert_to_pdf(output_path)
    
    print('\n分析完成！')

if __name__ == '__main__':
    main()
