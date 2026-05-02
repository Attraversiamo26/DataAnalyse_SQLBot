# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import numpy as np

# 尝试导入PDF转换库
try:
    from docx2pdf import convert as _docx2pdf_convert
except:
    _docx2pdf_convert = None

try:
    import win32com.client as _win32com
except:
    _win32com = None

import subprocess

# Set Chinese font
plt.rcParams['font.sans-serif'] = ['STSong']
plt.rcParams['axes.unicode_minus'] = False

# 使用相对路径配置
script_dir = os.path.dirname(os.path.abspath(__file__))
base_dir = os.path.dirname(script_dir)
xlsx_path_shared = os.path.join(base_dir, '..', 'data', '寄递时限对标分析指标大全.xlsx')
xlsx_path_local = os.path.join(base_dir, '输入数据', '接口文档-0401(1).xlsx')

if os.path.exists(xlsx_path_shared):
    xlsx_path = xlsx_path_shared
    print('使用共享数据文件：' + xlsx_path)
elif os.path.exists(xlsx_path_local):
    xlsx_path = xlsx_path_local
    print('使用本地数据文件：' + xlsx_path)
else:
    xlsx_path = xlsx_path_shared
    print('使用默认数据文件：' + xlsx_path)
output_dir = os.path.join(base_dir, '输出结果')
if not os.path.exists(output_dir):
    os.makedirs(output_dir, exist_ok=True)

# 读取各sheet数据
def read_sheets():
    # 线路表
    df_raw = pd.read_excel(xlsx_path, sheet_name='线路表', header=None)
    cols_route = df_raw.iloc[1].tolist()
    df_route = df_raw.iloc[2:].copy()
    df_route.columns = cols_route
    df_route.columns = [str(c).strip() for c in df_route.columns]
    
    # 投递机构表
    df_raw2 = pd.read_excel(xlsx_path, sheet_name='投递机构表', header=None)
    cols_delivery = df_raw2.iloc[1].tolist()
    df_delivery = df_raw2.iloc[2:].copy()
    df_delivery.columns = cols_delivery
    df_delivery.columns = [str(c).strip() for c in df_delivery.columns]
    
    # 48小时表
    df_48h = pd.read_excel(xlsx_path, sheet_name='进口、出口处理及投递48小时', header=0)
    df_48h['机构名称'] = df_48h['机构名称'].astype(str)
    df_48h['环节标识'] = df_48h['环节标识'].astype(str)
    df_48h['到达分拣离开标识'] = df_48h['到达分拣离开标识'].astype(str)
    df_48h['线路名'] = df_48h['线路名'].astype(str)
    
    return df_route, df_delivery, df_48h

def convert_to_pdf(docx_path):
    try:
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'

        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            chinese_font_paths = [
                '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
                'C:\\Windows\\Fonts\\simhei.ttf',
                '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf'
            ]
            for fp in chinese_font_paths:
                if os.path.exists(fp):
                    try:
                        pdf.add_font('chinese', '', fp)
                        pdf.set_font('chinese', '', 14)
                        break
                    except Exception:
                        continue
            else:
                pdf.set_font('Helvetica', '', 14)

            doc = Document(docx_path)
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    pdf.ln(5)
                    continue
                style_name = para.style.name if para.style else ''
                if 'Heading 1' in style_name:
                    pdf.set_font('chinese', '', 16)
                    pdf.ln(8)
                elif 'Heading 2' in style_name:
                    pdf.set_font('chinese', '', 14)
                    pdf.ln(5)
                elif 'Heading 3' in style_name:
                    pdf.set_font('chinese', '', 12)
                    pdf.ln(3)
                else:
                    pdf.set_font('chinese', '', 11)
                pdf.multi_cell(0, 8, text)
                pdf.ln(2)

            try:
                for table in doc.tables:
                    pdf.ln(3)
                    pdf.set_font('chinese', '', 8)
                    for row_idx, row in enumerate(table.rows):
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                pdf.multi_cell(0, 5, cell_text)
                        pdf.ln(2)
                    pdf.ln(2)
            except Exception:
                pass

            pdf.output(pdf_path)
            print('PDF文档已保存（fpdf2）: ' + pdf_path)
            return pdf_path
        except Exception as e:
            print('fpdf2转换失败: ' + str(e))

        if _docx2pdf_convert is not None:
            try:
                _docx2pdf_convert(docx_path, pdf_path)
                print('PDF文档已保存: ' + pdf_path)
                return pdf_path
            except Exception as e1:
                print('docx2pdf转换失败: ' + str(e1))

        if _win32com is not None:
            try:
                word = _win32com.client.Dispatch('Word.Application')
                word.Visible = False
                word_doc = word.Documents.Open(docx_path)
                word_doc.SaveAs(pdf_path, FileFormat=17)
                word_doc.Close()
                word.Quit()
                print('PDF文档已保存: ' + pdf_path)
                return pdf_path
            except Exception as e2:
                print('win32com转换失败: ' + str(e2))

        try:
            lo_paths = ['libreoffice', 'soffice',
                        os.path.expanduser('~/.libreoffice/LibreOffice.app/Contents/MacOS/soffice'),
                        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                        os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(base_dir))), '.libreoffice', 'LibreOffice.app', 'Contents', 'MacOS', 'soffice')]
            lo_cmd = None
            for p in lo_paths:
                if os.path.isabs(p):
                    if os.path.exists(p):
                        lo_cmd = p
                        break
                else:
                    from shutil import which
                    if which(p):
                        lo_cmd = p
                        break
            if not lo_cmd:
                raise FileNotFoundError('LibreOffice未安装')
            result = subprocess.run(
                [lo_cmd, '--headless', '--convert-to', 'pdf', '--outdir',
                 os.path.dirname(pdf_path), docx_path],
                capture_output=True, text=True, timeout=60
            )
            if os.path.exists(pdf_path):
                print('PDF文档已保存（LibreOffice）: ' + pdf_path)
                return pdf_path
        except FileNotFoundError:
            print('LibreOffice未安装，跳过PDF转换')
        except Exception as e3:
            print('LibreOffice转换异常: ' + str(e3))

        print('PDF转换失败，仅生成Word文档')
        return None
    except Exception as e:
        print('PDF转换失败: ' + str(e))
        return None

def create_48h_chart(df_48h, inst_name, output_dir):
    """生成48小时到达/下段/投递分布图"""
    inst_data = df_48h[(df_48h['机构名称'].str.contains(inst_name.replace(' ', ''), na=False)) & (df_48h.get('环节标识', '') == '投递')]
    if inst_data.empty:
        return None
    
    # 提取0-23小时的列
    hour_cols = []
    for c in df_48h.columns:
        try:
            if isinstance(c, (int, float)) and 0 <= float(c) <= 23:
                hour_cols.append(c)
        except:
            pass
    hour_cols = sorted(hour_cols, key=lambda x: int(float(x)))
    
    if not hour_cols:
        return None
    
    # 提取各标识的数据
    charts = {}
    for label in ['到达', '下段', '投递']:
        label_data = inst_data[inst_data['到达分拣离开标识'].str.contains(label, na=False)]
        if not label_data.empty:
            row = label_data.iloc[0]
            values = []
            for c in hour_cols:
                try:
                    val = float(row[c]) if pd.notna(row[c]) else 0
                    values.append(val)
                except:
                    values.append(0)
            charts[label] = values
    
    if not charts:
        return None
    
    # 绘制组合图
    fig, ax = plt.subplots(figsize=(12, 6))
    x = range(len(hour_cols))
    colors = {'到达': '#1f77b4', '下段': '#ff7f0e', '投递': '#2ca02c'}
    markers_map = {'到达': 'o', '下段': 's', '投递': '^'}
    
    for label, values in charts.items():
        color = colors.get(label, '#999999')
        marker = markers_map.get(label, 'o')
        ax.plot(x, values, marker=marker, label=label, color=color, linewidth=2, markersize=6)
    
    ax.set_xlabel('Hour', fontsize=12)
    ax.set_ylabel('Ratio', fontsize=12)
    ax.set_title(inst_name + ' - 48h Distribution', fontsize=14, fontweight='bold')
    ax.legend(prop={'size': 10})
    ax.set_xticks(x)
    ax.set_xticklabels([str(int(c)) for c in hour_cols], rotation=45)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, '48h_{}_综合分布图.png'.format(inst_name))
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    # 波峰分析
    peak_info = {}
    for label, values in charts.items():
        if values and max(values) > 0:
            max_idx = np.argmax(values)
            peak_info[label] = {'hour': str(int(hour_cols[max_idx])), 'value': values[max_idx] * 100}
    
    return chart_path, peak_info

def main():
    print('读取数据...')
    df_route, df_delivery, df_48h = read_sheets()
    
    # 创建Word文档
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('滨州市惠州市 线路投递环节时限对标下钻分析报告', 0)
    for run in title.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(24)
        run.font.bold = True
    
    date_str = datetime.now().strftime('%Y年%m月%d日')
    date_para = doc.add_paragraph('分析日期: ' + date_str)
    for run in date_para.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
    
    # 一、全程时限对标
    heading1 = doc.add_heading('一、全程时限对标', level=1)
    for run in heading1.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    # 查找滨州市惠州市线路数据
    route_data = df_route[df_route['线路名'] == '滨州市惠州市']
    if not route_data.empty:
        row = route_data.iloc[0]
        
        # 提取关键数据
        sample_post = pd.to_numeric(row.get('特快-样本量（件）', 0), errors='coerce')
        time_post = pd.to_numeric(row.get('特快-全程时限（小时）', 0), errors='coerce')
        time_comp = pd.to_numeric(row.get('竞品-全程时限（小时）', 0), errors='coerce')
        time_diff = pd.to_numeric(row.get('差值-全程时限（小时）', 0), errors='coerce')
        
        trajectory_post = str(row.get('主要轨迹-特快', ''))
        trajectory_comp = str(row.get('主要轨迹-竞品', ''))
        
        # 计算经转城市数量
        post_cities = len([c for c in trajectory_post.split('-') if c.strip()]) if trajectory_post else 0
        comp_cities = len([c for c in trajectory_comp.split('-') if c.strip()]) if trajectory_comp else 0
        
        # 投递环节时限
        delivery_post = pd.to_numeric(row.get('特快-投递时限（小时）', 0), errors='coerce')
        delivery_comp = pd.to_numeric(row.get('竞品-投递时限（小时）', 0), errors='coerce')
        delivery_diff = pd.to_numeric(row.get('差值-投递时限（小时）', 0), errors='coerce')
        
        # 生成分析文本
        if pd.notna(time_post) and pd.notna(time_comp):
            if time_post > time_comp:
                diff_text = '特快比竞品慢{:.1f}小时'.format(time_post - time_comp)
            else:
                diff_text = '特快比竞品快{:.1f}小时'.format(time_comp - time_post)
            
            if post_cities == comp_cities:
                traj_text = '特快及竞品主要轨迹相同，均为{}'.format(trajectory_post)
            elif post_cities > comp_cities:
                traj_text = '特快主要轨迹为{}，比竞品多经转{}个城市'.format(trajectory_post, post_cities - comp_cities)
            else:
                traj_text = '竞品主要轨迹为{}，比特快多经转{}个城市'.format(trajectory_comp, comp_cities - post_cities)
            
            full_para = doc.add_paragraph(
                '以滨州市至惠州市线路为例，{}，{}。{}'.format(
                    traj_text, diff_text,
                    '其中投递环节特快时限为{:.1f}小时，竞品为{:.1f}小时'.format(
                        delivery_post if pd.notna(delivery_post) else 0,
                        delivery_comp if pd.notna(delivery_comp) else 0) if pd.notna(delivery_post) and pd.notna(delivery_comp) else '')
            )
            for run in full_para.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(12)
            
            # 在终端打印
            print('\n' + '=' * 70)
            print('滨州市至惠州市 线路全程时限对标分析')
            print('=' * 70)
            print('样本量：{:.0f}件'.format(sample_post if pd.notna(sample_post) else 0))
            print('特快全程时限：{:.1f}小时'.format(time_post if pd.notna(time_post) else 0))
            print('竞品全程时限：{:.1f}小时'.format(time_comp if pd.notna(time_comp) else 0))
            print('全程时限差异：{}'.format(diff_text))
            if pd.notna(delivery_post) and pd.notna(delivery_comp):
                if delivery_post > delivery_comp:
                    delivery_text = '特快投递比竞品慢{:.1f}小时'.format(delivery_post - delivery_comp)
                else:
                    delivery_text = '特快投递比竞品快{:.1f}小时'.format(delivery_comp - delivery_post)
                print('投递环节差异：{}'.format(delivery_text))
    else:
        print('未找到滨州市惠州市线路数据')
        para = doc.add_paragraph('未找到滨州市至惠州市线路的完整数据。')
        for run in para.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
        return
    
    # 二、投递环节对标
    heading2 = doc.add_heading('二、投递环节对标', level=1)
    for run in heading2.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    if pd.notna(delivery_post) and pd.notna(delivery_comp):
        if delivery_post > delivery_comp:
            delivery_analysis = '投递环节特慢于竞品{:.1f}小时，是需要重点优化的环节。'.format(delivery_post - delivery_comp)
        else:
            delivery_analysis = '投递环节特快优于竞品{:.1f}小时，是邮政的竞争优势环节。'.format(delivery_comp - delivery_post)
        
        para2 = doc.add_paragraph(
            '投递环节方面，特快投递时限为{:.1f}小时，竞品为{:.1f}小时。{}'.format(
                delivery_post, delivery_comp, delivery_analysis)
        )
        for run in para2.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
        
        print('\n投递环节分析：{}'.format(delivery_analysis))
    
    # 三、投递机构下钻分析
    heading3 = doc.add_heading('三、投递机构下钻分析', level=1)
    for run in heading3.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    # 获取end5机构
    route_delivery = df_delivery[df_delivery['线路名'] == '滨州市惠州市'].copy()
    if not route_delivery.empty:
        for col in ['样本量（件）', '下段时间', '妥投时间（小时）', '投递时间（小时）', '多次投递占比']:
            if col in route_delivery.columns:
                route_delivery[col] = pd.to_numeric(route_delivery[col], errors='coerce')
        route_delivery = route_delivery.sort_values(by=['投递时间（小时）', '样本量（件）'], ascending=[False, False])
        end5 = route_delivery.head(5)
        
        # end5表格
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        headers = ['投递机构名称', '样本量（件）', '下段时间（小时）', '妥投时间（小时）', '投递时间（小时）', '多次投递占比']
        for j, h in enumerate(headers):
            hdr_cells[j].text = h
            for paragraph in hdr_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
                    run.font.bold = True
        
        for _, r in end5.iterrows():
            row_cells = table.add_row().cells
            for j, h in enumerate(headers):
                val = r.get(h, '')
                if pd.isna(val):
                    val = '-'
                elif isinstance(val, float):
                    if '占比' in h:
                        val = '{:.1f}%'.format(val * 100)
                    else:
                        val = '{:.2f}'.format(val)
                row_cells[j].text = str(val)
                for paragraph in row_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
        
        doc.add_paragraph('')
        
        # 在终端打印end5机构数据
        print('\n' + '-' * 70)
        print('投递时限end5问题机构：')
        print('-' * 70)
        print('{:<20} {:>8} {:>10} {:>10} {:>10} {:>10}'.format(
            '机构名称', '样本量', '下段时间', '妥投时间', '投递时间', '多次投递占比'))
        print('-' * 70)
        for _, r in end5.iterrows():
            name = str(r.get('投递机构名称', ''))[:18]
            sample = '{:.0f}'.format(r.get('样本量（件）', 0)) if pd.notna(r.get('样本量（件）', 0)) else '-'
            down = '{:.2f}'.format(r.get('下段时间', 0)) if pd.notna(r.get('下段时间', 0)) else '-'
            tuotou = '{:.2f}'.format(r.get('妥投时间（小时）', 0)) if pd.notna(r.get('妥投时间（小时）', 0)) else '-'
            delivery_val = '{:.2f}'.format(r.get('投递时间（小时）', 0)) if pd.notna(r.get('投递时间（小时）', 0)) else '-'
            multi = '{:.1f}%'.format(r.get('多次投递占比', 0) * 100) if pd.notna(r.get('多次投递占比', 0)) else '-'
            print('{:<20} {:>8} {:>10} {:>10} {:>10} {:>10}'.format(
                name, sample, down, tuotou, delivery_val, multi))
        
        # 四、问题机构48小时分布分析
        heading4 = doc.add_heading('四、问题机构48小时分布分析', level=1)
        for run in heading4.runs:
            run.font.name = '微软雅黑'
            run.font.bold = True
        
        # 针对最严重问题机构进行分析
        worst_inst = end5.iloc[0]
        worst_inst_name = str(worst_inst.get('投递机构名称', ''))
        
        if worst_inst_name:
            inst_heading = doc.add_heading('{} 下钻分析'.format(worst_inst_name), level=2)
            for run in inst_heading.runs:
                run.font.name = '微软雅黑'
                run.font.bold = True
            
            # 数据描述
            sample = worst_inst.get('样本量（件）', 0)
            down_time = worst_inst.get('下段时间', 0)
            tuotou_time = worst_inst.get('妥投时间（小时）', 0)
            delivery_time = worst_inst.get('投递时间（小时）', 0)
            multi_rate = worst_inst.get('多次投递占比', 0)
            
            detail_parts = []
            if pd.notna(sample):
                detail_parts.append('样本量{:.0f}件'.format(sample))
            if pd.notna(down_time):
                detail_parts.append('下段时间{:.2f}小时'.format(down_time))
            if pd.notna(tuotou_time):
                detail_parts.append('妥投时间{:.2f}小时'.format(tuotou_time))
            if pd.notna(delivery_time):
                detail_parts.append('投递时间{:.2f}小时'.format(delivery_time))
            if pd.notna(multi_rate):
                detail_parts.append('多次投递占比{:.1f}%'.format(multi_rate * 100))
            
            detail_para = doc.add_paragraph('该机构投递环节数据：' + '，'.join(detail_parts) + '。')
            for run in detail_para.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(12)
            
            # 生成48小时分布图
            chart_result = create_48h_chart(df_48h, worst_inst_name, output_dir)
            if chart_result:
                chart_path, peak_info = chart_result
                if os.path.exists(chart_path):
                    doc.add_paragraph('')
                    doc.add_paragraph('{} 48小时到达/下段/投递分布图：'.format(worst_inst_name))
                    doc.add_picture(chart_path, width=Inches(6))
                    
                    # 在终端显示分布数据表格
                    print('\n' + '=' * 70)
                    print('{} 48小时分布数据'.format(worst_inst_name))
                    print('=' * 70)
                    
                    inst_data_48h = df_48h[(df_48h['机构名称'].str.contains(worst_inst_name.replace(' ', ''), na=False)) & (df_48h.get('环节标识', '') == '投递')]
                    hour_cols_chart = sorted([c for c in df_48h.columns if isinstance(c, (int, float)) and 0 <= float(c) <= 23], key=lambda x: int(float(x)))
                    
                    charts_data = {}
                    for label in ['到达', '下段', '投递']:
                        label_data = inst_data_48h[inst_data_48h['到达分拣离开标识'].str.contains(label, na=False)]
                        if not label_data.empty:
                            row = label_data.iloc[0]
                            values = []
                            for c in hour_cols_chart:
                                try:
                                    val = float(row[c]) if pd.notna(row[c]) else 0
                                    values.append(val)
                                except:
                                    values.append(0)
                            charts_data[label] = values
                    
                    # 打印完整句子形式的分布数据
                    print('\n【各环节48小时分布详情】')
                    print('-' * 70)
                    
                    for label, values in charts_data.items():
                        print('\n{}环节分布情况：'.format(label))
                        
                        # 找出有数据的时间点
                        active_hours = [(int(h), values[int(h)]) for h in hour_cols_chart if int(h) < len(values) and values[int(h)] > 0]
                        
                        if not active_hours:
                            print('  该环节暂无数据。')
                            continue
                        
                        # 生成完整句子描述
                        if len(active_hours) == 1:
                            h, v = active_hours[0]
                            print('  所有邮件均集中在{}点{}，占比{:.1f}%，呈现高度集中的分布特征。'.format(
                                h, label, v * 100))
                        elif len(active_hours) == 2:
                            h1, v1 = active_hours[0]
                            h2, v2 = active_hours[1]
                            print('  邮件主要分布在{}点和{}点两个时间段，其中{}点占比{:.1f}%，{}点占比{:.1f}%。'.format(
                                h1, h2, h1, v1 * 100, h2, v2 * 100))
                            if v1 > v2:
                                print('  {}点为关键波峰时段，占比最高，达到{:.1f}%。'.format(h1, v1 * 100))
                            else:
                                print('  {}点为关键波峰时段，占比最高，达到{:.1f}%。'.format(h2, v2 * 100))
                        else:
                            # 找出波峰
                            peak_h, peak_v = max(active_hours, key=lambda x: x[1])
                            total_active = sum(v for _, v in active_hours)
                            main_hours = sorted([h for h, v in active_hours if v >= peak_v * 0.5])
                            
                            print('  该环节邮件分布在多个时间段，主要集中在{}点至{}点之间。'.format(
                                min(main_hours), max(main_hours)))
                            print('  关键波峰出现在{}点，占比{:.1f}%，是邮件到达/处理/投递的最集中时段。'.format(
                                peak_h, peak_v * 100))
                            
                            if peak_v > 0.5:
                                print('  该时段邮件量占比超过50%，说明存在明显的时段集中现象，可能导致处理能力不足。')
                    
                    print('\n' + '-' * 70)
                    
                    print('\n[图表已保存]: ' + chart_path)
                    
                    # 波峰分析
                    peak_parts = []
                    for label, info in peak_info.items():
                        peak_parts.append('{}关键波峰在{}点，占比{:.1f}%'.format(label, info['hour'], info['value']))
                    
                    analysis_para = doc.add_paragraph('波峰分析：' + '；'.join(peak_parts) + '。')
                    for run in analysis_para.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(12)
                    
                    # 问题诊断
                    diagnosis_heading = doc.add_heading('问题诊断', level=3)
                    for run in diagnosis_heading.runs:
                        run.font.name = '微软雅黑'
                        run.font.bold = True
                    
                    if '投递' in peak_info:
                        peak_hour = peak_info['投递']['hour']
                        peak_value = peak_info['投递']['value']
                        diag_para = doc.add_paragraph('{}投递时间长达{:.2f}小时，主要原因分析：'.format(worst_inst_name, delivery_time if pd.notna(delivery_time) else 0))
                        for run in diag_para.runs:
                            run.font.name = '微软雅黑'
                            run.font.size = Pt(12)
                        
                        diagnosis_parts = [
                            '{}%的投递集中在{}点完成，但妥投时间长达{:.2f}小时，说明投递过程耗时过长。'.format(peak_value, peak_hour, tuotou_time if pd.notna(tuotou_time) else 0),
                            '{:.1f}%的邮件需要多次投递，进一步拉长整体时限。'.format(multi_rate * 100 if pd.notna(multi_rate) else 0),
                            '下段时间较短（{:.2f}小时），说明问题不在下段环节，而在妥投环节。'.format(down_time if pd.notna(down_time) else 0)
                        ]
                        for part in diagnosis_parts:
                            p = doc.add_paragraph(part)
                            for run in p.runs:
                                run.font.name = '微软雅黑'
                                run.font.size = Pt(12)
                    
                    # 优化建议
                    suggestion_heading = doc.add_heading('优化建议', level=3)
                    for run in suggestion_heading.runs:
                        run.font.name = '微软雅黑'
                        run.font.bold = True
                    
                    suggestions = [
                        '优化投递路线，提高首次投递成功率，减少多次投递占比',
                        '加强客户联系确认，提前预约投递时间，避免无效投递',
                        '增加投递人员配置，提高单位时间投递效率',
                        '优化投递时间段安排，避开客户工作繁忙时段'
                    ]
                    for i, sug in enumerate(suggestions, 1):
                        p = doc.add_paragraph('{}. {}'.format(i, sug))
                        for run in p.runs:
                            run.font.name = '微软雅黑'
                            run.font.size = Pt(12)
    else:
        print('\n未找到滨州市惠州市线路的投递机构数据')
    
    # 五、分析总结
    heading5 = doc.add_heading('五、分析总结', level=1)
    for run in heading5.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True
    
    if pd.notna(time_post) and pd.notna(time_comp):
        summary_parts = [
            '滨州市至惠州市线路特快全程时限为{:.1f}小时，竞品为{:.1f}小时'.format(time_post, time_comp),
            '特快比竞品{}{:.1f}小时'.format('慢' if time_post > time_comp else '快', abs(time_post - time_comp))
        ]
        
        if pd.notna(delivery_post) and pd.notna(delivery_comp):
            summary_parts.append('投递环节特快为{:.1f}小时，竞品为{:.1f}小时'.format(delivery_post, delivery_comp))
            if delivery_post > delivery_comp:
                summary_parts.append('投递环节慢于竞品{:.1f}小时'.format(delivery_post - delivery_comp))
            else:
                summary_parts.append('投递环节优于竞品{:.1f}小时'.format(delivery_comp - delivery_post))
        
        summary_parts.append('需重点关注投递环节优化，提升整体时效表现。')
        
        summary = doc.add_paragraph('。'.join(summary_parts) + '。')
        for run in summary.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
    
    # 保存文档
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    output_path = os.path.join(output_dir, '线路投递环节时限分析报告_滨州市惠州市_{}.docx'.format(timestamp))
    doc.save(output_path)
    print('\nWord文档已保存: ' + output_path)
    
    # 显示48小时分布图路径
    chart_files = [f for f in os.listdir(output_dir) if f.startswith('48h_') and '滨州市惠州市' not in f and f.endswith('.png')]
    if chart_files:
        chart_path = os.path.join(output_dir, chart_files[0])
        print('\n48小时分布图: ' + chart_path)
    
    # 转换为PDF
    convert_to_pdf(output_path)
    
    print('\n分析完成！')

if __name__ == '__main__':
    main()
